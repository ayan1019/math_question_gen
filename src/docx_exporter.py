# src/docx_exporter.py
from docx import Document
from pathlib import Path
from question_model import Question

def export_questions_to_docx(title: str, description: str, questions: list[Question], out_path: str):
    """
    Export questions to a Word document (.docx)
    """
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(description)
    doc.add_paragraph("")

    for q in questions:
        doc.add_heading(f"Q{q.order}: {q.question}", level=2)
        doc.add_paragraph(f"Instruction: {q.instruction}")
        doc.add_paragraph(f"Difficulty: {q.difficulty}")

        for i, opt in enumerate(q.options):
            prefix = "[Correct] " if i == q.correct_index else ""
            doc.add_paragraph(f"{prefix}{opt}")

        doc.add_paragraph("Explanation:")
        doc.add_paragraph(q.explanation)

        doc.add_paragraph(f"Subject: {q.subject}")
        doc.add_paragraph(f"Unit: {q.unit}")
        doc.add_paragraph(f"Topic: {q.topic}")
        doc.add_paragraph(f"Marks: {q.plusmarks}")

        # Insert image if present
        if q.image_tag:
            img_path = Path("output") / f"{q.image_tag}.png"
            if img_path.exists():
                doc.add_picture(str(img_path))

        doc.add_page_break()

    doc.save(out_path)
    print(f"Saved Word file to {out_path}")
